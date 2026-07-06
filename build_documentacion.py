from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


OUT = r"C:\Users\MSI\OneDrive\Desktop\practica\documentacion_trabajo_autonomo.docx"


def set_font(run, name="Arial", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_par(doc, text="", style=None, align=None, space_after=6, space_before=0, line=1.15):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_font(r)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
styles["Normal"].font.size = Pt(11)

for name, size, color in [("Heading 1", 18, "136F63"), ("Heading 2", 14, "123D56"), ("Heading 3", 12, "123D56")]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
r = title.add_run("Documento del Trabajo Autónomo")
set_font(r, size=26, bold=True, color="0F172A")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(6)
r = subtitle.add_run("Sistema web ULEAM de capacidades especiales")
set_font(r, size=14, bold=False, color="64748B")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Portada | Informe de análisis, diseño e implementación")
set_font(r, size=10, color="64748B")

doc.add_paragraph()

cover = doc.add_table(rows=4, cols=2)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.autofit = False
widths = [Inches(2.0), Inches(4.8)]
for row in cover.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

labels = ["Carrera", "Asignatura", "Integrantes", "Fecha"]
values = [
    "No proporcionada",
    "Trabajo Autónomo / Desarrollo Web",
    "No proporcionados por el usuario",
    "06 de julio de 2026",
]
for idx, row in enumerate(cover.rows):
    row.cells[0].text = labels[idx]
    row.cells[1].text = values[idx]
    set_cell_shading(row.cells[0], "E7F8F4")
    set_cell_shading(row.cells[1], "FFFFFF")
    for j, cell in enumerate(row.cells):
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_font(run, bold=(j == 0), color="0F172A" if j else "0F4F47")

doc.add_paragraph()

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=18, bold=True, color="136F63")


def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color="123D56")


section_heading("1. Portada")
add_par(doc, "Este documento consolida el análisis, diseño e implementación del sistema web desarrollado en Vue 3 para la gestión de estudiantes, seguimientos, apoyos, adaptaciones, usuarios y reportes.", space_after=4)
add_par(doc, "Integrantes del grupo: No proporcionados por el usuario. Se recomienda reemplazar este dato antes de la entrega final.", space_after=4)

section_heading("2. Descripción del negocio o tema elegido")
add_par(doc, "El proyecto aborda la gestión de un sistema web institucional para el área de bienestar estudiantil y seguimiento académico de personas con capacidades especiales. El dominio cubre el registro y control de estudiantes, usuarios por rol, seguimientos académicos, apoyos, adaptaciones y reportes administrativos.", space_after=4)
add_par(doc, "El sistema centraliza información que antes podía estar dispersa en archivos manuales o formularios aislados, facilitando la consulta por parte de administradores, docentes y bienestar.", space_after=4)

section_heading("3. Propósito del sistema web")
add_par(doc, "El propósito principal del sistema es digitalizar y organizar el seguimiento integral de los estudiantes, permitiendo una administración más ordenada, rápida y consistente de la información.", space_after=4)
bullet("Registrar y consultar estudiantes con sus datos académicos y de discapacidad.")
bullet("Gestionar seguimientos académicos por docente y por estudiante.")
bullet("Registrar apoyos, observaciones y adaptaciones desde bienestar.")
bullet("Administrar usuarios, perfiles y roles del sistema.")
bullet("Generar reportes internos para toma de decisiones.")

section_heading("4. Requerimientos funcionales y no funcionales")
subheading("4.1 Requerimientos funcionales")
bullet("Permitir autenticación de usuarios por correo y contraseña.")
bullet("Gestionar estudiantes mediante operaciones CRUD.")
bullet("Permitir a administradores administrar usuarios, configuraciones y reportes.")
bullet("Permitir a docentes registrar seguimientos académicos de sus estudiantes asignados.")
bullet("Permitir a bienestar registrar apoyos, observaciones, adaptaciones y actualización de estado.")
bullet("Permitir búsquedas, filtros y visualización de detalle en cada módulo.")
bullet("Almacenar la información en el navegador usando localStorage.")

subheading("4.2 Requerimientos no funcionales")
bullet("La interfaz debe ser responsive y usable en pantallas de escritorio y móviles.")
bullet("La aplicación debe mantener consistencia visual entre módulos.")
bullet("El tiempo de respuesta debe ser inmediato al operar sobre datos locales.")
bullet("El sistema debe ser fácil de mantener gracias a la separación por componentes.")
bullet("La información debe persistir entre sesiones del navegador mientras no se borre el almacenamiento local.")

section_heading("5. Herramientas y tecnologías de programación")
add_par(doc, "El sistema fue desarrollado con un enfoque moderno de frontend, usando Vue 3 como framework principal y Vite como herramienta de desarrollo y construcción.", space_after=4)
bullet("Vue 3 para la construcción de componentes e interfaces.")
bullet("Vue Router para la navegación entre módulos.")
bullet("JavaScript para la lógica de negocio y validaciones.")
bullet("CSS personalizado para el diseño visual de la plataforma.")
bullet("localStorage para el almacenamiento de datos del lado del cliente.")
bullet("Font Awesome para iconografía de la interfaz.")
bullet("Vite para ejecución en desarrollo y build del proyecto.")
add_par(doc, "Los datos del sistema se guardan en estructuras serializadas tipo JSON dentro del almacenamiento local del navegador, lo que elimina la dependencia de una base de datos externa en esta versión académica.", space_after=4)

section_heading("6. Método de publicación y hosting")
add_par(doc, "Para la entrega académica, el sistema se ejecuta con Vite en entorno de desarrollo y se genera una compilación estática lista para publicación con build de producción.", space_after=4)
bullet("Modo de desarrollo: `npm run dev`.")
bullet("Build de producción: `npm run build`.")
bullet("Despliegue sugerido: hosting estático como Netlify, Vercel, GitHub Pages o servidor institucional.")
bullet("Persistencia: datos locales en localStorage del navegador.")
add_par(doc, "En esta documentación no se adjunta un enlace público de hosting porque no fue proporcionado por el usuario. Si existe, conviene agregar la URL de despliegue final en este apartado.", space_after=4)

section_heading("7. Código fuente completo")
add_par(doc, "El código fuente completo se encuentra organizado en el proyecto Vue dentro de la carpeta src, con módulos separados para admin, docente, bienestar, servicios y estilos globales.", space_after=4)
bullet("Carpeta principal: `src/`.")
bullet("Vistas: `src/views/`.")
bullet("Servicios: `src/services/`.")
bullet("Utilidades: `src/utils/`.")
bullet("Estilos: `src/assets/css/`.")
add_par(doc, "Enlace GitLab: No proporcionado por el usuario. Reemplazar este texto por la URL del repositorio final antes de entregar.", space_after=4)

section_heading("8. Manual de usuario del sistema web")
subheading("8.1 Acceso")
bullet("Abrir la aplicación desde el navegador.")
bullet("Ingresar correo institucional y contraseña.")
bullet("Seleccionar la opción de recordar sesión si está disponible.")

subheading("8.2 Módulo Administrador")
bullet("Desde el panel principal puede revisar estudiantes, usuarios, seguimientos y reportes.")
bullet("En usuarios puede crear, editar, activar o desactivar cuentas.")
bullet("En estudiantes puede consultar fichas, filtrar registros y editar información.")
bullet("En configuración puede ajustar parámetros generales y catálogos.")

subheading("8.3 Módulo Docente")
bullet("Visualiza sus estudiantes asignados y sus alertas académicas.")
bullet("Registra nuevos seguimientos con notas, asistencia y observaciones.")
bullet("Consulta historial por estudiante y filtra registros.")

subheading("8.4 Módulo Bienestar")
bullet("Consulta estudiantes con riesgo y su información integral.")
bullet("Registra apoyos, observaciones y adaptaciones.")
bullet("Actualiza el estado y nivel de riesgo del estudiante.")

subheading("8.5 Cierre de sesión")
bullet("Usar la opción Cerrar sesión en el menú lateral para salir de forma segura.")

section_heading("9. Conclusión del proyecto")
add_par(doc, "El sistema web desarrollado permite centralizar la gestión de estudiantes con capacidades especiales y mejorar el control académico y de bienestar mediante una interfaz modular, clara y funcional. La migración a Vue 3 fortaleció la mantenibilidad del proyecto y permitió una mejor organización por roles y vistas.", space_after=4)
add_par(doc, "Como cierre, el proyecto cumple con una solución de frontend moderna y escalable para un contexto académico, dejando abierta la posibilidad de integrar una base de datos y un backend en futuras versiones.", space_after=4)

section_heading("10. Bibliografía")
bullet("Vue.js Documentation. https://vuejs.org/")
bullet("Vue Router Documentation. https://router.vuejs.org/")
bullet("Vite Documentation. https://vite.dev/")
bullet("MDN Web Docs. https://developer.mozilla.org/")
bullet("OpenAI Codex documentation and project notes used during implementation.")

doc.save(OUT)
print(OUT)
